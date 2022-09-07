from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.firefox.options import Options
from bs4 import BeautifulSoup
import requests
from lxml import etree
import re
import time
import docx
from docx.enum.section import WD_ORIENT
from docx.shared import Cm
from Chapter import Chapter
import csv
import concurrent.futures
import os.path
from os import path


subjectURLDict = {
	'Botany': 'https://www.neetprep.com/questions/53-Botany?courseId=8',
	'Chemistry': 'https://www.neetprep.com/questions/54-Chemistry?courseId=8',
	'Physics': 'https://www.neetprep.com/questions/55-Physics?courseId=8',
	'Zoology': 'https://www.neetprep.com/questions/56-Zoology?courseId=8'
}

class Question:
	def __init__(self, qNo, text, a, b, c, d, e, correct):
		self.qNo = qNo
		self.text = text
		self.a = a
		self.b = b
		self.c = c
		self.d = d
		self.e = e
		self.correct = correct
		
# Lay BeautifulSoup cua website
def getSoup(site):
	options = Options()
	options.headless = True
	driver = webdriver.Firefox(options=options)
	# driver = webdriver.Firefox()
	# driver.maximize_window()
	driver.get(site)
	time.sleep(2)
	soup = BeautifulSoup(driver.page_source,'lxml')
	driver.quit()
	return soup

# Lay BeautifulSoup cua website và click tìm câu trả lời
def getSoupWithAnswer(site):
	options = Options()
	options.headless = True
	driver = webdriver.Firefox(options=options)
	# driver = webdriver.Firefox()
	# driver.maximize_window()
	driver.get(site)
	time.sleep(2)

	# Click để tìm câu trả lời đúng (correct answer)
	answerList = driver.find_elements(By.CLASS_NAME, 'option-list')
	for answer in answerList:
		answer.find_element(By.TAG_NAME, 'input').click()

	soup = BeautifulSoup(driver.page_source,'lxml')
	driver.quit()
	return soup


def mathlmToWord(mathml_string):
	tree = etree.fromstring(mathml_string)
	xslt = etree.parse('MML2OMML.XSL')
	transform = etree.XSLT(xslt)
	new_dom = transform(tree)
	return new_dom.getroot()

def htmlTableToText(tableSoup):
	result = ''
	rows = tableSoup.find_all('tr')
	for row in rows:
		cells = row.find_all('td')
		for cell in cells:
			for element in cell.find_all(text=True): #lấy tất cả element con có text thay bằng text
				element.replace_with(element.text)
			result += cell.text.strip()
			result += '_'
		result = result[:-1]
		result += '~'
	result = result[:-1]
	rows2 = result.split('~')
	if len(rows2) > 1:
		columns1 = rows2[0].split('_')
		columns2 = rows2[1].split('_')
		if len(columns1) != len(columns2):
			result = result[result.index('~')+1:]
	# print(result)
	return result

def writeElementListToCell(_elementList, cell):
	elementList = _elementList
	text = elementList[0]
	# print(question.text[0])
	elementList.pop(0)

	if '<math>' in text or '<img>' in text or '<table>' in text:
		stringList = re.split(r'<math>|<img>|<table>',text) # cắt chuỗi theo các cụm <math>|<img>|<table>
		
		# get previou string
		preString = '' 
		for string in stringList[:-1]:
			element = elementList[0] # lấy element đầu
			if '.png' in preString and ('=' not in string and (string == '' and 'math' not in element)): # nếu element trước đó là ảnh thì xuống đoạn khác
				cell.add_paragraph()
			cell.paragraphs[-1].add_run(string) # duyệt các string sau khi cắt chuối text
			if 'math' in element:
				try:
					cell.paragraphs[-1]._element.append(mathlmToWord(element))
				except:
					cell.paragraphs[-1].add_run(element)
					print('Error formula: ' + element)
					pass
			elif '.png' in element:
				try:
					cell.add_paragraph().add_run().add_picture(element, width = Cm(6.5))
				except:
					cell.add_paragraph().add_run(element)
					print('Error image: ' + element)
					pass
			elif '~' in element:
				textToDocTable(element, cell)

			preString = elementList.pop(0)

		if '.png' in preString and ('=' not in string and (string == '' and 'math' not in element)):
			cell.add_paragraph().add_run(stringList[-1])
		else:
			cell.paragraphs[-1].add_run(stringList[-1])
	else:
		cell.text = text

def textToDocTable(text, cell):
	rowsStr = text.split('~')
	# print(str(len(rowsStr)))
	# print(str(len(rowsStr[0].split('_'))))
	table = cell.add_table(rows=len(rowsStr), cols=len(rowsStr[0].split('_')))
	i = j = 0
	for row in rowsStr:
		cellsStr = row.split('_')
		for string in cellsStr:
			# print(str(i) + str(j) + string)
			table.rows[i].cells[j].text = string
			j += 1
		j = 0
		i += 1

def renumberQuestionList(questionList):
	newNo = int(questionList[0].qNo)

	# put available question into questionList2
	questionList2 = []
	for question in questionList:
		if question.text[0] == '' or question.a[0] == '' or question.b[0] == '' or question.c[0] == '' or question.d[0] == '':				
			# ghi log file
			with open('logfile.log', 'a') as log:
				log.write(question.qNo + '\n')
		else:
			# đổi số thứ tự câu hỏi
			question.qNo = str(newNo)
			questionList2.append(question)
			newNo += 1
	return questionList2


def writeDocFile(questionList, chapterName):
	doc = docx.Document()
	section = doc.sections[0]
	section.page_height = Cm(29.7)
	section.page_width = Cm(86)
	section.left_margin = Cm(2)
	section.right_margin = Cm(2)
	section.top_margin = Cm(2)
	section.bottom_margin = Cm(2)
	section.orientation = WD_ORIENT.LANDSCAPE
	# new_width, new_height = section.page_height, section.page_width
	# section.page_width = new_width
	# section.page_height = new_height
	table = doc.add_table(rows=1, cols=16)
	for i in range(16):
		if i == 0:
			table.columns[i].width = Cm(2)
		elif i == 1:
			table.columns[i].width = Cm(10)
		else:
			table.columns[i].width = Cm(5)
	row = table.rows[0].cells
	row[0].text = 'Q No'
	row[1].text = 'Qtext'
	row[2].text = 'A'
	row[3].text = 'B'
	row[4].text = 'C'
	row[5].text = 'D'
	row[6].text = 'E'
	row[7].text = 'Correct'
	row[8].text = 'Exp'
	row[9].text = 'Hint'
	row[10].text = 'Level'
	row[11].text = 'Mark'
	row[12].text = 'Time Duration'
	row[13].text = 'Chapter No'
	row[14].text = 'Topic No'
	row[15].text = 'Paragraph Text'

	for question in questionList:
		row = table.add_row().cells
		row[0].text = question.qNo
		print(question.qNo)

		# put question text into column 1
		writeElementListToCell(question.text, row[1])

		# put question a into column 2
		writeElementListToCell(question.a, row[2])

		# put question b into column 3
		writeElementListToCell(question.b, row[3])

		# put question c into column 4
		writeElementListToCell(question.c, row[4])

		# put question d into column 5
		writeElementListToCell(question.d, row[5])

		# put question e into column 6
		writeElementListToCell(question.e, row[6])

		row[7].text = question.correct
		row[8].text = ''
		row[9].text = ''
		row[10].text = ''
		row[11].text = ''
		row[12].text = ''
		row[13].text = ''
		row[14].text = ''
		row[15].text = ''
	
	doc.save(chapterName + '.docx')


def getQuestionsOnePage(pageURL, chapterName):
	print(pageURL)
	# soup = getSoupWithAnswer(pageURL)
	soup = getSoup(pageURL)

	questionList = []

	questionListHtml = soup.find_all('div', class_='question-body')
	for questionHtml in questionListHtml:
		q = Question('', [''], [''], [''], [''], [''], [''], '')

		# get Question No
		q.qNo = questionHtml.find('div', class_='question-tag').text.replace('Q', '').replace(':', '').strip()
		print(q.qNo)

		questionTextHtml = questionHtml.find('div', class_='question-text').find('span')

		# scrap <br>
		for br in questionTextHtml('br'):
			br.replace_with('\n')

		# get Question text and table

		# get textHtml:
		textHtml = questionTextHtml.find_all(recursive=False)


		if len(textHtml) == 0:
			allText = questionTextHtml.text.replace('\xa0', '')
			stringList = re.split(r'\n *1.|\n *\(1\)|\n *2.|\n *\(2\)|\n *3.|\n *\(3\)|\n *4.|\n *\(4\)|\n *5.|\n *\(5\)',allText)
			# print(stringList[0])
			q.text[0] = stringList[0]
			q.a[0] = stringList[1]
			q.b[0] = stringList[2]
			q.c[0] = stringList[3]
			q.d[0] = stringList[4]
			if len(stringList) > 5:
				q.e[0] = stringList[5]

		else:
			imageCount = 1

			for element in textHtml: # mỗi element là 1 thẻ trong question-text

				questionTextList = []

				# get table
				if element.name == 'table':
					questionTextList.append(htmlTableToText(element))

				# get p tag
				elif element.name == 'p':

					# get all direct child of p tag (diagram, mathjax)
					childElements = element.find_all(recursive=False)

					# if there is any child
					if len(childElements) > 0:
						for child in childElements:
							if child.name == 'span' and child.has_attr('class') and child['class'][0] == 'mjx-chtml': # con là span và class MathJax
								questionTextList.append(child['data-mathml'])		# nối chuỗi MathJax vào questionTextList
								child.replace_with('<math>')	# thay tag span thành <math>

							elif child.name == 'img': # img
								try:
									img_data = requests.get(child['src'], timeout=2).content
									with open('Images/' + chapterName + '_' + q.qNo + '_' + str(imageCount) + '.png', 'wb') as handler: # ghi file ảnh
										handler.write(img_data)
									questionTextList.append('Images/' + chapterName + '_' + q.qNo + '_' + str(imageCount) + '.png') # nối path của ảnh vào questionTextList
									child.replace_with('<img>')	# thay tag img thành <img>
									imageCount += 1
								except:
									child.decompose()
									print('Error image: ' + q.qNo + '_' + chapterName)
									questionTextList = ['error']
									break
									
							elif child.name == 'b' or child.name == 'i' or (child.name == 'span' and child.text is not None): 
								child.replace_with(child.text)

							else: 
								child.decompose()

					# insert p text into first of questionTextList
					if element.text != '':
						questionTextList.insert(0, element.text.replace('\xa0', ''))

				# questiontextlist is empty
				if len(questionTextList) == 0:
					continue

				# put questionText into question text or question options
				try:
					if 'error' in questionTextList:
						q.text = ['']
						break
					elif '1.' in questionTextList[0] and questionTextList[0].index('1.') == 0:
						questionTextList[0] = questionTextList[0][2:].strip()
						q.a = questionTextList
					elif '(1)' in questionTextList[0] and questionTextList[0].index('(1)') == 0:
						questionTextList[0] = questionTextList[0][3:].strip()
						q.a = questionTextList
					elif '2.' in questionTextList[0] and questionTextList[0].index('2.') == 0:
						questionTextList[0] = questionTextList[0][2:].strip()
						q.b = questionTextList
					elif '(2)' in questionTextList[0] and questionTextList[0].index('(2)') == 0:
						questionTextList[0] = questionTextList[0][3:].strip()
						q.b = questionTextList
					elif '3.' in questionTextList[0] and questionTextList[0].index('3.') == 0:
						questionTextList[0] = questionTextList[0][2:].strip()
						q.c = questionTextList
					elif '(3)' in questionTextList[0] and questionTextList[0].index('(3)') == 0:
						questionTextList[0] = questionTextList[0][3:].strip()
						q.c = questionTextList
					elif '4.' in questionTextList[0] and questionTextList[0].index('4.') == 0:
						questionTextList[0] = questionTextList[0][2:].strip()
						q.d = questionTextList
					elif '(4)' in questionTextList[0] and questionTextList[0].index('(4)') == 0:
						questionTextList[0] = questionTextList[0][3:].strip()
						q.d = questionTextList
					elif '5.' in questionTextList[0] and questionTextList[0].index('5.') == 0:
						questionTextList[0] = questionTextList[0][2:].strip()
						q.e = questionTextList
					elif '(5)' in questionTextList[0] and questionTextList[0].index('(5)') == 0:
						questionTextList[0] = questionTextList[0][3:].strip()
						q.e = questionTextList
					elif '|' in questionTextList[0] and '_' in questionTextList[0]: # nếu là bảng và qText không rỗng
						if q.text[0] != '':
							q.text[0] += '<table>'
							q.text.append(questionTextList[0])
							if '|2._' in questionTextList[0] or '|(2)_' in questionTextList[0] : # nếu là bảng trả lời
								q.a = ['1']
								q.b = ['2']
								q.c = ['3']
								q.d = ['4']
								q.e = ['5']
						else:
							continue
					else:
						q.text[0] += questionTextList[0]
						q.text.extend(questionTextList[1:])	# nối thêm questionTextList với các p textList mới
				except:
					# create q.text with empty string in list
					q.text = ['']

					print('Error qText: ' + q.qNo + '_' + chapterName)
					break

		# get correct answer
		if len(questionHtml.find_all('div', class_='_2eaw _2kqr')) > 0:
			q.correct = questionHtml.find_all('div', class_='_2eaw _2kqr')[0].find_all('label')[0].text
		
		if '1' in q.correct:
			q.correct = 'A'
		elif '2' in q.correct:
			q.correct = 'B'
		elif '3' in q.correct:
			q.correct = 'C'
		elif '4' in q.correct:
			q.correct = 'D'
		elif '5' in q.correct:
			q.correct = 'E'

		questionList.append(q)

	return questionList

def getQuestionsOfChapter(chapter):

	questionList = []
	numberOfPages = int(int(chapter.numberOfQuestions)/10) + 1

	for page in range(numberOfPages):
		questionList.extend(getQuestionsOnePage(chapter.chapterURL + '&pageNo=' + str(page+1), chapter.chapterName))

	writeDocFile(renumberQuestionList(questionList), chapter.chapterName)

#----------------------------------------

if __name__ == '__main__':

	chapterList = []

	with open('Chapters.csv') as file:
		reader = csv.reader(file)
		for row in reader:
			chapter = Chapter(*row)
			if chapter.subjectName == 'Zoology':
				# print(chapter.chapterName)
				chapterList.append(chapter)

	# get question of particular chapter
	getQuestionsOfChapter(chapterList[0])

	# create list from particular chapters
	# chapterList = [chapterList[1], chapterList[2], chapterList[4]]

	# run multithread to get from list of chapter
	# with concurrent.futures.ThreadPoolExecutor() as executor:
	# 	results = executor.map(getQuestionsOfChapter, chapterList)

	# for chapter in chapterList: # check error chapter when finish
	# 	if path.exists(chapter.chapterName + '.docx') is False:
	# 		print('Error chapter: ' + chapter.chapterName + '_' + str(chapterList.index(chapter)))

	# run loop all page in 1 chapter to find out error page
	# for i in range(27):
	# 	questionList = getQuestionsOnePage('https://www.neetprep.com/questions/54-Chemistry/647-Classification-Elements-Periodicity-Properties?courseId=8&pageNo=' + str(i+1), 'DemoData')
	# 	writeDocFile(renumberQuestionList(questionList), 'DemoData')

	# get question in 1 particular page
	# questionList = getQuestionsOnePage('https://www.neetprep.com/questions/54-Chemistry/647-Classification-Elements-Periodicity-Properties?courseId=8&pageNo=27', 'DemoData')
	# writeDocFile(renumberQuestionList(questionList), 'DemoData')