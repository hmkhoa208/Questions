from docx import Document
from lxml import etree

# Convert MathML (MML) into Office MathML (OMML) using a XSLT stylesheet

mathml_string = '<math xmlns="http://www.w3.org/1998/Math/MathML"><mn>15</mn><mfenced close="|" open="|"><mrow><mn>5</mn><mover><mi mathvariant="normal">i</mi><mo>^</mo></mover><mo> </mo><mo>+</mo><mo> </mo><mn>6</mn><mover><mi mathvariant="normal">j</mi><mo>^</mo></mover><mo> </mo><mo>+</mo><mo> </mo><mn>5</mn><mover><mi mathvariant="normal">k</mi><mo>^</mo></mover></mrow></mfenced></math>'

def mathlmToWord(mathml_string):
    tree = etree.fromstring(mathml_string)
    xslt = etree.parse('MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    return new_dom.getroot()

doc = Document()
paragraph = doc.add_paragraph()
paragraph.add_run('Function: ')
paragraph._element.append(mathlmToWord(mathml_string))
paragraph.add_run(' end')
doc.add_paragraph().add_run().add_picture('Q10:.png', width = 1000000, height = 1000000)
doc.save('test.docx')
